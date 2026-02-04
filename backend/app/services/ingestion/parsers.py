import os
import statistics
from typing import List
from collections import Counter

import fitz  # PyMuPDF
import pdfplumber
from docx import Document

from app.services.ingestion.interfaces import (
    DocumentParser,
    ParsedDocument,
    ParsedBlock,
    ParsedSpan,
    ParsedTable,
)


class PDFParser(DocumentParser):
    """
    PDF parser using PyMuPDF for text + structure and pdfplumber for tables.

    Two-pass approach:
    1. PyMuPDF: Extract text with font metadata (size, bold, italic) for heading detection.
    2. pdfplumber: Extract tables as structured data with markdown representation.
    """

    def parse(self, file_path: str) -> str:
        """Backward-compatible: returns plain text."""
        doc = self.parse_structured(file_path)
        return doc.full_text

    def parse_structured(self, file_path: str) -> ParsedDocument:
        """Parse PDF into structured document with headings, blocks, and tables."""
        filename = os.path.basename(file_path)

        # Pass 1: PyMuPDF for text with structure
        blocks, all_spans, total_pages, page_dimensions = self._extract_text_with_structure(file_path)

        # Detect heading thresholds from font statistics
        body_size = self._detect_body_font_size(all_spans)
        self._assign_heading_levels(blocks, body_size)

        # Compute median font size for slide detection
        median_font_size = 0.0
        if all_spans:
            font_sizes = [s.font_size for s in all_spans if s.font_size > 0]
            if font_sizes:
                median_font_size = statistics.median(font_sizes)

        # Pass 2: pdfplumber for tables
        tables = self._extract_tables(file_path)

        # Build full text and headings list
        full_text_parts = []
        headings = []
        for block in blocks:
            full_text_parts.append(block.content)
            if block.heading_level > 0:
                headings.append(block.content.strip())

        for table in tables:
            if table.markdown:
                full_text_parts.append(table.markdown)

        full_text = "\n\n".join(full_text_parts)

        return ParsedDocument(
            filename=filename,
            total_pages=total_pages,
            blocks=blocks,
            tables=tables,
            full_text=full_text,
            headings=headings,
            page_dimensions=page_dimensions,
            median_font_size=median_font_size,
        )

    def _extract_text_with_structure(self, file_path: str):
        """Extract text blocks with font metadata using PyMuPDF."""
        doc = fitz.open(file_path)
        blocks = []
        all_spans = []
        total_pages = len(doc)

        # Capture page dimensions from first page for slide detection
        page_dimensions = ()
        if total_pages > 0:
            rect = doc[0].rect
            page_dimensions = (rect.width, rect.height)

        for page_num in range(total_pages):
            page = doc[page_num]
            page_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)

            for block in page_dict.get("blocks", []):
                if block.get("type") != 0:  # Skip non-text blocks (images)
                    continue

                block_text_parts = []
                block_spans = []

                for line in block.get("lines", []):
                    line_text_parts = []
                    for span in line.get("spans", []):
                        text = span.get("text", "").strip()
                        if not text:
                            continue

                        font_size = span.get("size", 12)
                        flags = span.get("flags", 0)
                        is_bold = bool(flags & 2**4)  # bit 4 = bold
                        is_italic = bool(flags & 2**1)  # bit 1 = italic

                        parsed_span = ParsedSpan(
                            text=text,
                            font_size=font_size,
                            is_bold=is_bold,
                            is_italic=is_italic,
                            bbox=tuple(span.get("bbox", ())),
                        )
                        block_spans.append(parsed_span)
                        all_spans.append(parsed_span)
                        line_text_parts.append(text)

                    if line_text_parts:
                        block_text_parts.append(" ".join(line_text_parts))

                block_content = "\n".join(block_text_parts).strip()
                if not block_content:
                    continue

                blocks.append(
                    ParsedBlock(
                        page_number=page_num + 1,
                        block_type="text",
                        content=block_content,
                        heading_level=0,  # Will be assigned later
                        spans=block_spans,
                        bbox=tuple(block.get("bbox", ())),
                    )
                )

        doc.close()
        return blocks, all_spans, total_pages, page_dimensions

    def _detect_body_font_size(self, spans: List[ParsedSpan]) -> float:
        """Detect the most common (body) font size using mode."""
        if not spans:
            return 12.0

        sizes = [round(s.font_size, 1) for s in spans if s.font_size > 0]
        if not sizes:
            return 12.0

        # Use mode (most frequent size) as the body text size
        size_counts = Counter(sizes)
        body_size = size_counts.most_common(1)[0][0]
        return body_size

    def _assign_heading_levels(self, blocks: List[ParsedBlock], body_size: float):
        """Assign heading levels based on font size relative to body text."""
        for block in blocks:
            if not block.spans:
                continue

            # Use the max font size in the block
            max_font_size = max(s.font_size for s in block.spans)
            # Check if most spans are bold
            bold_count = sum(1 for s in block.spans if s.is_bold)
            is_mostly_bold = bold_count > len(block.spans) / 2

            # Short text blocks are more likely headings
            is_short = len(block.content) < 120

            if max_font_size >= body_size * 1.5 and is_short:
                block.heading_level = 1
            elif max_font_size >= body_size * 1.2 and is_short:
                block.heading_level = 2
            elif is_mostly_bold and is_short and max_font_size >= body_size:
                block.heading_level = 3

    def _extract_tables(self, file_path: str) -> List[ParsedTable]:
        """Extract tables using pdfplumber."""
        tables = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    page_tables = page.extract_tables()
                    if not page_tables:
                        continue

                    for raw_table in page_tables:
                        if not raw_table or len(raw_table) < 2:
                            continue

                        # Clean cells
                        cleaned_rows = []
                        for row in raw_table:
                            cleaned_row = [
                                (cell.strip() if cell else "") for cell in row
                            ]
                            # Skip completely empty rows
                            if any(c for c in cleaned_row):
                                cleaned_rows.append(cleaned_row)

                        if len(cleaned_rows) < 2:
                            continue

                        # First row is headers
                        headers = cleaned_rows[0]
                        data_rows = cleaned_rows[1:]

                        # Build markdown
                        markdown = self._table_to_markdown(headers, data_rows)

                        tables.append(
                            ParsedTable(
                                page_number=page_num,
                                rows=cleaned_rows,
                                headers=headers,
                                markdown=markdown,
                                caption="",
                            )
                        )
        except Exception as e:
            print(f"pdfplumber table extraction failed: {e}")

        return tables

    def _table_to_markdown(self, headers: List[str], data_rows: List[List[str]]) -> str:
        """Convert table to markdown format."""
        if not headers:
            return ""

        header_line = "| " + " | ".join(headers) + " |"
        separator = "| " + " | ".join(["---"] * len(headers)) + " |"

        lines = [header_line, separator]
        for row in data_rows:
            # Pad row if needed
            padded = row + [""] * (len(headers) - len(row))
            lines.append("| " + " | ".join(padded[: len(headers)]) + " |")

        return "\n".join(lines)


class DocxParser(DocumentParser):
    def parse(self, file_path: str) -> str:
        text = ""
        try:
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        except Exception as e:
            print(f"Error parsing DOCX {file_path}: {e}")
            raise
        return text


class DocumentParserFactory:
    @staticmethod
    def get_parser(file_path: str) -> DocumentParser:
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()

        if ext == ".pdf":
            return PDFParser()
        elif ext == ".docx":
            return DocxParser()
        else:
            raise ValueError(f"Unsupported file extension: {ext}")
